import PyPDF2
import pandas as pd
from pathlib import Path
from docx import Document
from dataclasses import dataclass
from typing import List, Dict, Any

from .config import Config

@dataclass
class DocumentChunk:
    """Represents a chunk of text from a document"""
    text: str
    metadata: Dict[str, Any]
    chunk_id: str


class PDFLoader:
    """Load and extract text from PDF files (technical papers in this case)"""
    
    @staticmethod
    def load(file_path: Path) -> List[Dict[str, Any]]:
        """
        Load PDF and extract text with page numbers
        
        Returns:
            List of dicts with 'text', 'page', 'source'
        """
        print(f"Loading PDF: {file_path.name}")
        
        chunks = []
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                total_pages = len(pdf_reader.pages)
                
                print(f"Total pages: {total_pages}")
                
                for page_num in range(total_pages):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    
                    # Clean text
                    text = text.strip()
                    
                    # Skip empty pages
                    if not text or len(text) < 50:
                        continue
                    
                    chunks.append({
                        'text': text,
                        'page': page_num + 1,
                        'source': file_path.name,
                        'doc_type': 'pdf'
                    })
                
                print(f"Extracted {len(chunks)} pages with content")
                return chunks
                
        except Exception as e:
            print(f"Error loading PDF: {e}")
            raise


class DOCXLoader:
    """Load and extract text from DOCX files (legal documents in this case)"""
    
    @staticmethod
    def load(file_path: Path) -> List[Dict[str, Any]]:
        """
        Load DOCX and extract text by paragraphs
        
        Returns:
            List of dicts with 'text', 'paragraph_num', 'source'
        """
        print(f"Loading DOCX: {file_path.name}")
        
        chunks = []
        
        try:
            doc = Document(file_path)
            total_paragraphs = len(doc.paragraphs)
            
            print(f"Total paragraphs: {total_paragraphs}")
            
            current_section = []
            section_num = 0
            
            for para_num, paragraph in enumerate(doc.paragraphs, 1):
                text = paragraph.text.strip()
                
                # Skip empty paragraphs
                if not text:
                    continue
                
                # Check if this is a section header (all caps or ends with colon)
                is_header = (text.isupper() or text.endswith(':')) and len(text) < 100
                
                if is_header and current_section:
                    # Save previous section
                    section_text = '\n'.join(current_section)
                    if len(section_text) > 100:  # Only save substantial sections
                        chunks.append({
                            'text': section_text,
                            'section': section_num,
                            'source': file_path.name,
                            'doc_type': 'docx'
                        })
                    current_section = [text]
                    section_num += 1
                else:
                    current_section.append(text)
            
            # Don't forget the last section
            if current_section:
                section_text = '\n'.join(current_section)
                if len(section_text) > 100:
                    chunks.append({
                        'text': section_text,
                        'section': section_num,
                        'source': file_path.name,
                        'doc_type': 'docx'
                    })
            
            print(f" Extracted {len(chunks)} sections")
            return chunks
            
        except Exception as e:
            print(f" Error loading DOCX: {e}")
            raise


class ExcelLoader:
    """Load and convert Excel to text descriptions (tabular data)"""
    
    @staticmethod
    def load(file_path: Path) -> List[Dict[str, Any]]:
        """
        Load Excel and create comprehensive text descriptions
        handles inflation calculator with years in first column
        
        Returns:
            List of dicts with 'text', 'sheet', 'row', 'decade', 'source', 'doc_type'
        """
        print(f"Loading Excel: {file_path.name}")
        
        chunks = []
        
        try:
            excel_file = pd.ExcelFile(file_path)
            
            print(f"Sheets found: {len(excel_file.sheet_names)}")
            
            for sheet_name in excel_file.sheet_names:
                # Read the sheet
                raw = pd.read_excel(file_path, sheet_name=sheet_name, header=None, engine="openpyxl", )

                # Find the cell that is exactly "Year"
                header_pos = None
                for r in range(raw.shape[0]):
                    for c in range(raw.shape[1]):
                        val = str(raw.iat[r, c]).strip()
                        if val == "Year":
                            header_pos = (r, c)
                            # do not break early if you want the last match; if you prefer last, remove these breaks
                            break
                    if header_pos:
                        break

                if header_pos:
                    header_row, header_col = header_pos

                    # Data starts on the next row after the header, from the Year column onwards
                    data = raw.iloc[header_row + 1 :, header_col :].copy()

                    # Set column names from the header row
                    data.columns = raw.iloc[header_row, header_col :].tolist()

                    df = data
                else:
                    # Fallback
                    df = pd.read_excel(
                        file_path,
                        sheet_name=sheet_name,
                        engine="openpyxl",
                    )

                
                print(f"Processing sheet '{sheet_name}': {len(df)} rows, {len(df.columns)} columns")
                print("Columns:", df.columns.tolist())
                
                # Identify the year column (first column)
                year_column = df.columns[0]
                print("year column", year_column)
                
                # Get actual year values
                years = df[year_column].dropna().tolist()
                if years:
                    start_year = years[0]
                    end_year = years[-1]
                else:
                    start_year = "unknown"
                    end_year = "unknown"
                
                # Create summary description
                columns = df.columns.tolist()
                summary = f"Sheet: {sheet_name}\n"
                summary += f"This is an inflation calculator with historical CPI data from {start_year} to {end_year}.\n"
                summary += f"Available data columns: Year, Jan, Feb, Mar, Apr, May, Jun, Jul, Aug, Sep, Oct, Nov, Dec, Average\n"
                summary += "This data can be used to calculate inflation-adjusted values between any two years.\n"
                
                chunks.append({
                    'text': summary,
                    'sheet': sheet_name,
                    'row': 'summary',
                    'source': file_path.name,
                    'doc_type': 'excel'
                })
                
                # Create comprehensive year-by-year entries
                for idx, row in df.iterrows():
                    # Get year from first column
                    year = row[year_column]
                    
                    # Skip if year is NaN or not a valid number
                    if pd.isna(year):
                        continue
                    
                    # Convert to int if it's a valid year
                    try:
                        year = int(float(year))
                    except (ValueError, TypeError):
                        continue
                    
                    # Skip header rows or invalid years
                    if year < 1900 or year > 2100:
                        continue
                    
                    # Create comprehensive description for this year
                    year_desc = f"Inflation data for year {year}:\n"
                    
                    # Add monthly values with proper column names
                    months = ['Jan', 'Feb', 'Mar', 'Apr', 'May', 'Jun', 
                             'Jul', 'Aug', 'Sep', 'Oct', 'Nov', 'Dec']
                    
                    monthly_data = []
                    for month in months:
                        # Find column that contains this month (case-insensitive)
                        month_col = None
                        for col in df.columns:
                            if str(col).strip().lower() == month.lower():
                                month_col = col
                                break
                        
                        if month_col is not None and pd.notna(row[month_col]):
                            value = row[month_col]
                            monthly_data.append(f"{month}: {value}")
                    
                    if monthly_data:
                        year_desc += "Monthly CPI values: " + ", ".join(monthly_data) + "\n"
                    
                    # Add average - look for 'Average' or 'Avg' column
                    avg_col = None
                    for col in df.columns:
                        col_str = str(col).strip().lower()
                        if 'average' in col_str or 'avg' in col_str:
                            avg_col = col
                            break
                    
                    if avg_col is not None and pd.notna(row[avg_col]):
                        avg_value = row[avg_col]
                        year_desc += f"Average annual CPI for {year}: {avg_value}\n"
                    
                    # Add contextual information
                    year_desc += f"Year {year} inflation data can be used to calculate inflation-adjusted values.\n"
                    year_desc += f"To adjust a value from {year} to another year, use the CPI ratio between those years."
                    
                    chunks.append({
                        'text': year_desc,
                        'sheet': sheet_name,
                        'row': int(idx) + 2,  # Excel row number
                        'year': year,
                        'source': file_path.name,
                        'doc_type': 'excel'
                    })
                
                # Create decade summaries for easier retrieval
                decades = {}
                for idx, row in df.iterrows():
                    year = row[year_column]
                    
                    if pd.isna(year):
                        continue
                    
                    try:
                        year = int(float(year))
                        if year < 1900 or year > 2100:
                            continue
                        
                        decade = (year // 10) * 10
                        if decade not in decades:
                            decades[decade] = []
                        
                        # Get average for this year
                        avg_col = None
                        for col in df.columns:
                            col_str = str(col).strip().lower()
                            if 'average' in col_str or 'avg' in col_str:
                                avg_col = col
                                break
                        
                        if avg_col is not None and pd.notna(row[avg_col]):
                            decades[decade].append((year, row[avg_col]))
                    except (ValueError, TypeError):
                        continue
                
                # Create decade summary chunks
                for decade, years_data in decades.items():
                    if years_data:
                        decade_desc = f"Inflation data summary for the {decade}s:\n"
                        decade_desc += f"Years covered: {min(y[0] for y in years_data)} to {max(y[0] for y in years_data)}\n"
                        decade_desc += "Year-by-year average CPI:\n"
                        
                        for year, avg in sorted(years_data):
                            decade_desc += f"  {year}: {avg}\n"
                        
                        chunks.append({
                            'text': decade_desc,
                            'sheet': sheet_name,
                            'row': f'decade_{decade}',
                            'decade': decade,
                            'source': file_path.name,
                            'doc_type': 'excel'
                        })
            
            print(f"Created {len(chunks)} searchable entries")
            return chunks
            
        except Exception as e:
            print(f"Error loading Excel: {e}")
            import traceback
            traceback.print_exc()
            raise

class DocumentLoader:
    """Main document loader that handles all file types"""
    
    def __init__(self):
        self.loaders = {
            '.pdf': PDFLoader(),
            '.docx': DOCXLoader(),
            '.xlsx': ExcelLoader()
        }
    
    def load_all_documents(self) -> Dict[str, List[Dict[str, Any]]]:
        """
        Load all challenge documents
        
        Returns:
            Dict mapping document names to their extracted chunks
        """
        print("\n" + "="*60)
        print("LOADING ALL DOCUMENTS")
        print("="*60 + "\n")
        
        Config.validate()  # Ensure all files exist
        
        doc_paths = Config.get_document_paths()
        results = {}
        total_chunks = 0
        
        for doc_name, doc_path in doc_paths.items():
            try:
                # Get appropriate loader based on file extension
                loader = self.loaders.get(doc_path.suffix.lower())
                
                if not loader:
                    print(f"No loader for {doc_path.suffix} files, skipping {doc_name}")
                    continue
                
                # Load document
                chunks = loader.load(doc_path)
                results[doc_name] = chunks
                total_chunks += len(chunks)
                
            except Exception as e:
                print(f"Failed to load {doc_name}: {e}")
                results[doc_name] = []
        
        print("\n" + "="*60)
        print(f"LOADING COMPLETE: {total_chunks} total chunks from {len(results)} documents")
        print("="*60 + "\n")
        
        return results
    
    def get_loading_summary(self, results: Dict[str, List[Dict[str, Any]]]) -> str:
        """Get a summary of loaded documents"""
        summary = "\nDocument Loading Summary:\n" + "="*50 + "\n"
        
        for doc_name, chunks in results.items():
            if chunks:
                doc_type = chunks[0].get('doc_type', 'unknown')
                summary += f"{doc_name} ({doc_type}): {len(chunks)} chunks\n"
                
                # Show sample metadata
                if chunks:
                    sample = chunks[0]
                    meta_keys = [k for k in sample.keys() if k != 'text']
                    summary += f"Metadata: {', '.join(meta_keys)}\n"
            else:
                summary += f"{doc_name}: Failed to load\n"
        
        summary += "="*50
        return summary
